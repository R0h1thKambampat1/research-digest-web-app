from flask import Flask, render_template, request, send_file
import feedparser
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.opc import constants
import os
import re
from datetime import datetime
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/scrape', methods=['POST'])
def scrape():
    query = request.form['query']

    heading = query
    query = "+".join(query.split())
    feed_url = f'https://export.arxiv.org/api/query?search_query=all:{query}&start=0&max_results=100'
    feed = feedparser.parse(feed_url)

    entries = sorted(feed.entries, key=lambda entry: datetime.strptime(entry.published, '%Y-%m-%dT%H:%M:%SZ'), reverse=True)

    doc = Document()
    doc.add_heading(f'ArXiv Papers for "{heading}"', level=1)

    for i, entry in enumerate(entries, 1):
        doc.add_heading(f'{i}. {entry.title}', level=1)
        authors = ', '.join(author.name for author in entry.authors)
        cleaned_summary = clean_latex(entry.summary)

        doc.add_paragraph(f"Authors: {authors}")
        doc.add_paragraph(f"Published: {entry.published}")
        doc.add_heading(f'Abstract', level=2)
        doc.add_paragraph(cleaned_summary)
        p = doc.add_paragraph()
        add_hyperlink(p, entry.link, "View on arXiv")
        doc.add_paragraph("-" * 50)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'{"_".join(heading.split(" "))}_{timestamp}.docx'
    filepath = os.path.join(os.getcwd(), filename)
    doc.save(filepath)
    response = send_file(filepath, as_attachment=True)
    os.remove(filepath)

    return response

def clean_latex(text):
    # Replace common LaTeX math symbols with Unicode
    latex_to_unicode = {
        r'\alpha': 'α',
        r'\beta': 'β',
        r'\gamma': 'γ',
        r'\delta': 'δ',
        r'\epsilon': 'ε',
        r'\pi': 'π',
        r'\mu': 'μ',
        r'\sigma': 'σ',
        r'\lambda': 'λ',
        r'\rightarrow': '→',
        r'\infty': '∞',
        r'\approx': '≈',
        r'\leq': '≤',
        r'\geq': '≥'
    }
    for latex, uni in latex_to_unicode.items():
        text = text.replace(latex, uni)

    # Remove math mode symbols like $, \[ \], \( \)
    text = re.sub(r'\$+', '', text)
    text = re.sub(r'\\\[|\\\]|\\\(|\\\)', '', text)

    # Remove LaTeX commands like \cite{}, \ref{}
    text = re.sub(r'\\(cite|ref|eqref|emph|textbf|textit)\{.*?\}', '', text)

    # Remove remaining LaTeX commands
    text = re.sub(r'\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^\}]*\})?', '', text)

    # Remove double spaces
    text = re.sub(r'\s{2,}', ' ', text)

    return text.strip()

def add_hyperlink(paragraph, url, text):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed attributes
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a w:rPr element (run properties)
    rpr = OxmlElement('w:rPr')

    # Style the hyperlink (blue and underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0000FF")
    rpr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rpr.append(underline)

    new_run.append(rpr)

    # Add the text to the run
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

if __name__ == '__main__':
    app.run(debug=True)
